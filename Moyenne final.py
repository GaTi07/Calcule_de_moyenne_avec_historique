# Créé par PEREIRA, le 06/10/2023 en Python 3.12

import subprocess

# Liste des bibliothèques à installer
libraries_to_install = ["openpyxl", "python-docx"]

# Fonction pour installer les bibliothèques
def install_libraries():
    for library in libraries_to_install:
        try:
            subprocess.check_call(["pip", "install", library])
        except subprocess.CalledProcessError as e:
            print(f"Erreur lors de l'installation de {library}: {e}")

# Vérifier si les bibliothèques sont installées
try:
    import openpyxl
    from docx import Document
except ImportError:
    print("Certaines bibliothèques ne sont pas installées. Installation en cours...")
    install_libraries()

# Importer les bibliothèques après installation
import openpyxl
from docx import Document

import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
from tkinter import filedialog
import openpyxl
from docx import Document

# Liste pour stocker l'historique des calculs et des moyennes
historique = []

# Palette de couleurs personnalisée pour un thème sombre
palette = {
    'background': '#8C8C8B',       # Fond Principal
    'text_primary': '#404040',      # Texte Principal
    'text_secondary': '#777777',    # Texte Secondaire
    'accent_color': '#0077b6',      # Couleur d'Accent
    'button_link_color': '#48cae4', # Boutons et Liens
    'button_bg': '#404040',         # Fond des boutons
    'button_fg': '#D9D7D7'          # Texte des boutons (Noir)
}

font_bold = ("Calibri", 19, "bold")  # Police en gras
font_normal = ("Calibri", 19)        # Police normale

expression_calcul = ""

# Fonction pour calculer la moyenne
def calculer_moyenne():
    global expression_calcul
    valeurs = entry.get()

    if not valeurs:
        messagebox.showerror("Erreur", "Veuillez entrer au moins un nombre.")
        return

    valeurs = valeurs.split(',')

    try:
        valeurs = [float(val.strip()) for val in valeurs]
        somme = sum(valeurs)
        nombre_de_valeurs = len(valeurs)
        moyenne = somme / nombre_de_valeurs
        expression_calcul = f"({'+'.join(map(str, valeurs))}) / {nombre_de_valeurs}"

        resultat_label.config(text=f"Moyenne : {moyenne:.2f}")

        # Ajoutez le calcul à la colonne "Action" et la moyenne à la colonne "Valeur" de l'historique
        historique.append((expression_calcul, f"{moyenne:.2f}"))
        afficher_historique()

    except ValueError:
        messagebox.showerror("Erreur", "Veuillez entrer des nombres valides.")

# Fonction pour afficher l'historique des calculs et des moyennes dans le tableau
def afficher_historique():
    historique_tree.delete(*historique_tree.get_children())
    for i, (action, valeur) in enumerate(historique, start=1):
        historique_tree.insert("", "end", values=(action, valeur))

# Fonction pour enregistrer tous les logs dans un fichier texte
def enregistrer_tout_logs():
    try:
        fichier = filedialog.asksaveasfile(defaultextension=".txt", filetypes=[("Fichiers texte", "*.txt")])

        if fichier:
            # En-tête du tableau
            fichier.write("Action\tValeur\n")

            for i, (action, valeur) in enumerate(historique, start=1):
                # Écrivez chaque ligne du tableau avec des onglets pour séparer les colonnes
                fichier.write(f"{action}\t{valeur}\n")

            messagebox.showinfo("Succès", "Tous les logs enregistrés avec succès dans le fichier.")
            fichier.close()
    except Exception as e:
        messagebox.showerror("Erreur", f"Une erreur s'est produite lors de l'enregistrement des logs : {str(e)}")

# Fonction pour enregistrer tous les logs dans un fichier Word
def enregistrer_logs_word():
    try:
        fichier = filedialog.asksaveasfile(defaultextension=".docx", filetypes=[("Fichiers Word", "*.docx")])

        if fichier:
            doc = Document()
            doc.add_heading("Historique des calculs et des moyennes", level=1)

            # Créez un tableau dans le document Word
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.allow_autofit = False

            # Ajoutez les en-têtes de colonne
            cellules_en_tetes = table.rows[0].cells
            cellules_en_tetes[0].text = "Calcul(s)"
            cellules_en_tetes[1].text = "Moyenne(s)"

            # Ajoutez les données de l'historique au tableau
            for action, valeur in historique:
                ligne = table.add_row().cells
                ligne[0].text = action
                ligne[1].text = valeur

            doc.save(fichier.name)
            messagebox.showinfo("Succès", "Les logs ont été enregistrés avec succès dans le fichier Word.")
    except Exception as e:
        messagebox.showerror("Erreur", f"Une erreur s'est produite lors de l'enregistrement des logs : {str(e)}")

# Fonction pour enregistrer tous les logs dans un fichier Excel
def enregistrer_logs_excel():
    try:
        fichier = filedialog.asksaveasfile(defaultextension=".xlsx", filetypes=[("Fichiers Excel", "*.xlsx")])

        if fichier:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Historique des calculs"

            # Ajoutez les en-têtes de colonne
            ws['A1'] = "Calcul(s)"
            ws['B1'] = "Moyenne(s)"

            # Ajoutez les données de l'historique au tableau
            for i, (action, valeur) in enumerate(historique, start=2):
                ws[f'A{i}'] = action
                ws[f'B{i}'] = float(valeur)

            wb.save(fichier.name)
            messagebox.showinfo("Succès", "Les logs ont été enregistrés avec succès dans le fichier Excel.")
    except Exception as e:
        messagebox.showerror("Erreur", f"Une erreur s'est produite lors de l'enregistrement des logs : {str(e)}")

# Fonction pour afficher les crédits
def afficher_credits():
    global fenetre
    fenetre.withdraw()  # Masquer la fenêtre principale

    # Créez une nouvelle fenêtre pour les crédits
    fenetre_credits = tk.Toplevel()
    fenetre_credits.title("Crédits")

    largeur_fenetre_credits = 400
    hauteur_fenetre_credits = 200

    fenetre_credits.geometry(f"{largeur_fenetre_credits}x{hauteur_fenetre_credits}")

    cadre_credits = tk.Frame(fenetre_credits, bg=palette['background'])
    cadre_credits.pack(fill="both", expand=True)

    # Ajoutez les informations de crédits
    credits_label = tk.Label(cadre_credits, text="Crédits", font=font_bold, bg=palette['background'], fg=palette['text_primary'])
    credits_label.pack(pady=20)

    auteur_label = tk.Label(cadre_credits, text="Auteur : PEREIRA TIAGO", font=font_bold, bg=palette['background'], fg=palette['text_primary'])
    auteur_label.pack()

    # Bouton pour revenir en arrière avec les mêmes couleurs que les autres boutons
    retour_button = tk.Button(cadre_credits, text="Retour", command=lambda: retourner_fenetre(fenetre_credits), font=font_bold, bg=palette['button_bg'], fg=palette['button_fg'], borderwidth=3)
    retour_button.pack(pady=10)

    fenetre_credits.mainloop()

# Fonction pour revenir à la fenêtre principale depuis la fenêtre des crédits
def retourner_fenetre(fenetre_a_afficher):
    global fenetre
    fenetre_a_afficher.destroy()  # Fermer la fenêtre actuelle
    fenetre.deiconify()  # Réafficher la fenêtre principale

# Création de la fenêtre principale
fenetre = tk.Tk()
fenetre.title("Calcul de moyenne avec historique")

largeur_fenetre = 600
hauteur_fenetre = 850  # Ajustement de la hauteur minimale

fenetre.geometry(f"{largeur_fenetre}x{hauteur_fenetre}")
fenetre.minsize(width=largeur_fenetre, height=hauteur_fenetre)  # Définir la taille minimale

cadre_principal = tk.Frame(fenetre, bg=palette['background'])
cadre_principal.pack(fill="both", expand=True)

instructions_label = tk.Label(cadre_principal, text="Entrez les nombres séparés par des virgules :", font=font_bold, bg=palette['background'], fg=palette['text_primary'])
instructions_label.pack(pady=10)

entry = tk.Entry(cadre_principal, font=font_normal, borderwidth=3)
entry.pack(pady=10)

# Créez un cadre pour les boutons
cadre_boutons = tk.Frame(cadre_principal, bg=palette['background'])
cadre_boutons.pack(fill="both", expand=True)

calculer_button = tk.Button(
    cadre_boutons,
    text="Calculer la moyenne",
    command=calculer_moyenne,
    font=font_bold,
    bg='#00FF00',  # Couleur verte pour ce bouton
    fg='#FFFFFF',  # Texte blanc
    borderwidth=3
)

calculer_button.pack(padx=10, pady=10)

# Créez un cadre pour afficher le résultat de la moyenne
cadre_resultat = tk.Frame(cadre_principal, bg=palette['background'])
cadre_resultat.pack(pady=10)

resultat_label = tk.Label(cadre_resultat, text="", font=font_bold, bg=palette['background'], fg=palette['text_primary'], borderwidth=3)
resultat_label.pack()

# Créez un cadre pour le tableau d'historique
cadre_historique = tk.Frame(cadre_principal, bg=palette['background'])
cadre_historique.pack(pady=10, fill="both", expand=True)

# Créez une étiquette pour l'historique des calculs et des moyennes
historique_label = tk.Label(cadre_historique, text="Historique des calculs et des moyennes :", font=font_bold, bg=palette['background'], fg=palette['text_primary'])
historique_label.pack()

# Créez un tableau (Treeview) pour afficher l'historique des calculs et des moyennes
historique_tree = ttk.Treeview(cadre_historique, columns=("Action", "Valeur"), show="headings")
historique_tree.heading("Action", text="Action")
historique_tree.heading("Valeur", text="Valeur")
historique_tree.pack(fill="both", expand=True)

# Créez un cadre pour les boutons Enregistrer tous les Logs en Word et Excel
cadre_boutons3 = tk.Frame(cadre_principal, bg=palette['background'])
cadre_boutons3.pack(fill="both", expand=True)

enregistrer_tout_logs_button = tk.Button(
    cadre_boutons3,
    text="Enregistrer logs en fichier Texte",
    command=enregistrer_tout_logs,
    font=font_bold,
    bg=palette['button_bg'],  # Fond blanc
    fg=palette['button_fg'],  # Texte noir
    borderwidth=3,
)
enregistrer_tout_logs_button.pack(padx=10, pady=10)

enregistrer_logs_word_button = tk.Button(
    cadre_boutons3,
    text="Enregistrer logs en Word",
    command=enregistrer_logs_word,
    font=font_bold,
    bg=palette['button_bg'],  # Fond blanc
    fg=palette['button_fg'],  # Texte noir
    borderwidth=3,
)
enregistrer_logs_word_button.pack(padx=10, pady=10)

enregistrer_logs_excel_button = tk.Button(
    cadre_boutons3,
    text="Enregistrer logs en Excel",
    command=enregistrer_logs_excel,
    font=font_bold,
    bg=palette['button_bg'],  # Fond blanc
    fg=palette['button_fg'],  # Texte noir
    borderwidth=3,
)
enregistrer_logs_excel_button.pack(padx=10, pady=10)

# Créez un cadre pour les boutons Crédits
cadre_boutons4 = tk.Frame(cadre_principal, bg=palette['background'])
cadre_boutons4.pack(fill="both", expand=True)

credits_button = tk.Button(
    cadre_boutons4,
    text="Crédits",
    command=afficher_credits,
    font=font_bold,
    bg=palette['button_bg'],  # Fond blanc
    fg=palette['button_fg'],  # Texte noir
    borderwidth=3,
)
credits_button.pack(padx=10, pady=10)

# Lancer la boucle principale
fenetre.mainloop()
